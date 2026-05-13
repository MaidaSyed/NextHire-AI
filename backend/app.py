import difflib
import io
import json
import logging
import os
import pathlib
import re
import string
import tempfile
import urllib.error
import urllib.request

from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from jinja2 import Environment, FileSystemLoader, select_autoescape

from config import get_setting

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    from dotenv import load_dotenv

    load_dotenv(dotenv_path=pathlib.Path(__file__).resolve().parent / ".env")
except Exception:
    pass

app = Flask(__name__)

# Enable CORS for development and production
CORS(app, 
     resources={r"/api/*": {"origins": "*", "methods": ["GET", "POST", "OPTIONS"], "allow_headers": ["Content-Type"]}},
     origins=["http://localhost:5173", "http://localhost:5174", "http://127.0.0.1:5173", "http://127.0.0.1:5174", "*"],
     supports_credentials=True)

# Production settings
app.config['JSON_SORT_KEYS'] = False
app.config['ENV'] = os.environ.get('FLASK_ENV', 'production')


_BASE_STOPWORDS = {
    "a",
    "an",
    "the",
    "and",
    "or",
    "of",
    "to",
    "in",
    "for",
    "with",
    "is",
    "are",
    "was",
    "were",
    "be",
    "been",
    "being",
    "on",
    "at",
    "by",
    "from",
    "as",
    "that",
    "this",
    "these",
    "those",
    "it",
    "its",
    "i",
    "we",
    "you",
    "they",
    "them",
    "our",
    "your",
    "their",
    "he",
    "she",
    "his",
    "her",
    "not",
    "but",
    "if",
    "then",
    "so",
    "such",
    "can",
    "could",
    "should",
    "would",
    "will",
    "may",
    "might",
    "must",
    "do",
    "does",
    "did",
    "done",
    "have",
    "has",
    "had",
    "having",
    "about",
    "into",
    "over",
    "under",
    "between",
    "while",
    "during",
    "before",
    "after",
    "up",
    "down",
    "out",
    "off",
    "again",
    "further",
    "here",
    "there",
    "when",
    "where",
    "why",
    "how",
    "all",
    "any",
    "both",
    "each",
    "few",
    "more",
    "most",
    "other",
    "some",
    "no",
    "nor",
    "only",
    "own",
    "same",
    "than",
    "too",
    "very",
}

_SYNONYM_TO_CANONICAL = {
    "teamwork": "collaboration",
    "collaboration": "collaboration",
    "collaborate": "collaboration",
    "api": "rest api",
    "apis": "rest api",
    "rest api": "rest api",
    "rest apis": "rest api",
    "data analysis": "analytics",
    "analytics": "analytics",
}

_PHRASE_SKILLS = {
    "data analysis",
    "machine learning",
    "deep learning",
    "natural language processing",
    "nlp",
    "computer vision",
    "rest api",
    "unit testing",
    "test automation",
    "cloud computing",
    "microservices",
    "project management",
    "agile methodology",
    "scrum master",
    "full stack",
    "frontend developer",
    "backend developer",
    "data science",
    "artificial intelligence",
}

_COMMON_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "nextjs",
    "vue",
    "angular",
    "node",
    "flask",
    "django",
    "fastapi",
    "sql",
    "postgresql",
    "mysql",
    "mongodb",
    "redis",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "git",
    "linux",
    "html",
    "css",
    "tailwind",
    "bootstrap",
    "graphql",
    "rest api",
    "microservices",
    "unit testing",
    "ci",
    "cd",
    "devops",
    "jenkins",
    "terraform",
    "ansible",
}


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        raise RuntimeError("PyPDF2 not installed. Run: python -m pip install PyPDF2")

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read PDF: {exc}")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("python-docx not installed. Run: python -m pip install python-docx")

    try:
        bio = io.BytesIO(file_bytes)
        doc = Document(bio)
        parts: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        # also try to extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX: {exc}")


def _get_stopwords() -> set[str]:
    try:
        import nltk
        from nltk.corpus import stopwords

        try:
            words = set(stopwords.words("english"))
        except Exception:
            nltk.download("stopwords", quiet=True)
            words = set(stopwords.words("english"))
        return _BASE_STOPWORDS | words
    except Exception:
        return set(_BASE_STOPWORDS)


def _tokenize(text: str, stopwords_set: set[str]) -> list[str]:
    if not text:
        return []

    lowered = str(text).lower()
    cleaned = re.sub(r"[^a-z\s]+", " ", lowered)
    cleaned = cleaned.translate(str.maketrans("", "", string.punctuation))
    tokens = cleaned.split()

    result: list[str] = []
    for t in tokens:
        if t in stopwords_set:
            continue
        if len(t) < 2:
            continue
        result.append(t)
    return result


def _normalize_term(term: str) -> str:
    t = term.strip().lower()
    if t in _SYNONYM_TO_CANONICAL:
        return _SYNONYM_TO_CANONICAL[t]
    if t.endswith("s") and len(t) > 3:
        singular = t[:-1]
        if singular in _SYNONYM_TO_CANONICAL:
            return _SYNONYM_TO_CANONICAL[singular]
    return t


def _extract_keywords(tokens: list[str]) -> set[str]:
    if not tokens:
        return set()

    out: set[str] = set()
    out.update(_normalize_term(t) for t in tokens)

    for i in range(len(tokens) - 1):
        phrase = f"{tokens[i]} {tokens[i+1]}"
        normalized = _normalize_term(phrase)
        if normalized in _PHRASE_SKILLS:
            out.add(normalized)

    return out


def _cosine_similarity_fallback(a_keywords: set[str], b_keywords: set[str]) -> float:
    if not a_keywords or not b_keywords:
        return 0.0
    
    # Standard intersection
    intersection = len(a_keywords & b_keywords)
    
    # Add fuzzy matching for similar words (e.g., "developer" vs "development")
    fuzzy_matches = 0
    a_list = list(a_keywords)
    b_list = list(b_keywords)
    
    # Only do fuzzy matching for keywords that aren't already in intersection
    remaining_a = [w for w in a_list if w not in b_keywords]
    remaining_b = [w for w in b_list if w not in a_keywords]
    
    for wa in remaining_a:
        for wb in remaining_b:
            # Simple prefix matching or very close similarity
            if wa.startswith(wb) or wb.startswith(wa) or difflib.SequenceMatcher(None, wa, wb).ratio() > 0.8:
                fuzzy_matches += 0.5
                break # Count each word at most once
                
    total_match = intersection + fuzzy_matches
    denom = (len(a_keywords) * len(b_keywords)) ** 0.5
    return min(1.0, total_match / denom) if denom else 0.0


def _semantic_similarity_score(resume_text: str, jd_text: str) -> float:
    if not resume_text or not jd_text:
        return 0.0

    try:
        from sentence_transformers import SentenceTransformer, util

        if not hasattr(_semantic_similarity_score, "_model"):
            _semantic_similarity_score._model = SentenceTransformer("all-MiniLM-L6-v2")
        model = _semantic_similarity_score._model

        emb1 = model.encode(resume_text, convert_to_tensor=True, normalize_embeddings=True)
        emb2 = model.encode(jd_text, convert_to_tensor=True, normalize_embeddings=True)
        sim = util.cos_sim(emb1, emb2).item()
        return max(0.0, min(1.0, float(sim))) * 100
    except Exception:
        pass

    # Fallback to keyword-based similarity if sentence-transformers fails
    stopwords_set = _get_stopwords()
    resume_keywords = _extract_keywords(_tokenize(resume_text, stopwords_set))
    jd_keywords = _extract_keywords(_tokenize(jd_text, stopwords_set))
    return max(0.0, min(1.0, _cosine_similarity_fallback(resume_keywords, jd_keywords))) * 100


def _analyze_formatting(text: str) -> list[str]:
    issues = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    
    if not lines:
        return ["Empty document detected."]
        
    # Check for long paragraphs (poor ATS readability)
    long_paragraphs = 0
    for line in lines:
        if len(line.split()) > 60:
            long_paragraphs += 1
    
    if long_paragraphs > 1:
        issues.append("Detected large text blocks - break them into shorter bullets for better ATS parsing.")
        
    # Check for bullet points (ATS likes them)
    bullet_points = [l for l in lines if l.startswith(('-', '•', '*', '+', '·'))]
    if len(bullet_points) < 5:
        issues.append("Low use of bullet points - use them for better scanability of your achievements.")
        
    # Check for common sections
    text_lower = text.lower()
    missing_sections = []
    if "experience" not in text_lower and "work" not in text_lower and "employment" not in text_lower:
        missing_sections.append("Experience")
    if "education" not in text_lower:
        missing_sections.append("Education")
    if "skill" not in text_lower:
        missing_sections.append("Skills")
        
    if missing_sections:
        issues.append(f"Standard sections missing or poorly labeled: {', '.join(missing_sections)}.")
        
    return issues


def calculate_advanced_ats(resume_text: str, jd_text: str, resume_source: str, jd_source: str) -> dict:
    stopwords_set = _get_stopwords()

    resume_tokens = _tokenize(resume_text, stopwords_set)
    jd_tokens = _tokenize(jd_text, stopwords_set)

    resume_keywords = _extract_keywords(resume_tokens)
    jd_keywords = _extract_keywords(jd_tokens)

    total_jd_keywords = len(jd_keywords)
    total_resume_keywords = len(resume_keywords)

    matched_keywords = sorted(jd_keywords & resume_keywords)
    missing_keywords = sorted(jd_keywords - resume_keywords)

    # Calculate keyword matching score
    keyword_score = 0.0
    if total_jd_keywords:
        keyword_score = (len(matched_keywords) / total_jd_keywords) * 100

    # Calculate semantic similarity
    semantic_score = _semantic_similarity_score(resume_text, jd_text)

    # Extract skills from JD and resume
    jd_skills = set()
    for k in jd_keywords:
        if k in _COMMON_SKILLS or k in _PHRASE_SKILLS:
            jd_skills.add(k)

    resume_skills = set()
    for k in resume_keywords:
        if k in _COMMON_SKILLS or k in _PHRASE_SKILLS:
            resume_skills.add(k)

    important_skills_detected = sorted(resume_skills & jd_skills) if jd_skills else sorted(resume_skills)

    # Calculate skill matching score
    skill_score = 0.0
    if jd_skills:
        skill_score = (len(resume_skills & jd_skills) / len(jd_skills)) * 100
    else:
        skill_score = keyword_score

    # Analyze formatting early to use in scoring
    formatting_issues = _analyze_formatting(resume_text)

    # Enhanced scoring: Weighted formula
    # 40% for core skill matching, 30% semantic relevance, 30% keyword/experience matching
    # Add a small base score if there is any relevance at all
    base_score = 15.0 if (skill_score > 0 or semantic_score > 20 or keyword_score > 0) else 0.0
    
    raw_ats_score = (0.4 * skill_score) + (0.3 * semantic_score) + (0.3 * keyword_score)
    ats_score = base_score + (raw_ats_score * 0.85) # Scale down to fit base score
    
    # Penalize for formatting issues (max -10 points)
    if formatting_issues:
        penalty = min(10, len(formatting_issues) * 2.5)
        ats_score = max(0, ats_score - penalty)

    # Boost score if resume has good semantic match (indicates strong contextual fit)
    if semantic_score > 60 and skill_score > 50:
        ats_score = min(98, ats_score + 7)
    
    # Ensure the score reflects reality - if semantic and skills are both high, final score should be high
    if semantic_score > 70 and skill_score > 60:
        ats_score = max(ats_score, 75)
    
    # Cap at 100
    ats_score = min(100, ats_score)

    suggestions: list[str] = []
    
    # Intelligent suggestions based on gaps
    if skill_score < 50:
        missing_top_skills = list(missing_keywords)[:5]
        if missing_top_skills:
            suggestions.append(f"Add key skills: {', '.join(missing_top_skills)}")
    
    if semantic_score < 60:
        suggestions.append("Tailor experience descriptions to match job responsibilities")
    elif semantic_score < 75:
        suggestions.append("Enhance project descriptions with specific achievements and metrics")
    
    if len(resume_tokens) < max(150, len(jd_tokens) // 2):
        suggestions.append("Expand experience and project descriptions")
    
    if keyword_score < 50 and missing_keywords:
        suggestions.append("Include more industry-specific terminology from the job description")
    
    if not suggestions:
        suggestions.append("Strong match - consider applying!")

    seen = set()
    deduped_suggestions: list[str] = []
    for s in suggestions:
        if s not in seen:
            deduped_suggestions.append(s)
            seen.add(s)

    return {
        "ats_score": round(float(ats_score), 2),
        "keyword_score": round(float(keyword_score), 2),
        "semantic_score": round(float(semantic_score), 2),
        "skill_score": round(float(skill_score), 2),
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "important_skills_detected": important_skills_detected,
        "formatting_issues": formatting_issues,
        "resume_source": resume_source,
        "jd_source": jd_source,
        "suggestions": deduped_suggestions,
    }


def _error(message: str, status_code: int = 400):
    return jsonify({"error": message}), status_code


def _backend_templates_root() -> pathlib.Path:
    return pathlib.Path(__file__).resolve().parent / "templates"


def _jinja_env() -> Environment:
    if not hasattr(_jinja_env, "_env"):
        _jinja_env._env = Environment(
            loader=FileSystemLoader(str(_backend_templates_root())),
            autoescape=select_autoescape(["html", "xml"]),
        )
    return _jinja_env._env


def _template_exists(template_id: str) -> bool:
    t_root = _backend_templates_root() / template_id
    return (t_root / "index.html").is_file() and (t_root / "style.css").is_file()


# Prepended to every resume stylesheet for WeasyPrint / wkhtmltopdf: A4 page box,
# no extra body padding (otherwise 210mm + padding exceeds the sheet and clips),
# flex children must be allowed to shrink, long URLs must wrap.
_RESUME_RENDER_BASE_CSS = """
@page { size: A4; margin: 0; }
html, body {
  margin: 0 !important;
  padding: 0 !important;
  height: auto !important;
  max-width: 100%;
}
.page {
  box-sizing: border-box;
  max-width: 100%;
}
.sidebar, .main-content, .left-column, .right-column, .left-sidebar, .right-content,
.container, .content, .header {
  min-width: 0;
}
.entry-header, .experience-header, .education-header {
  min-width: 0;
}
.entry-title, .job-title-main, .degree-name, .exp-title, .edu-degree,
.summary-text, .section-text, .contact-info, .contact-link, .detail-item,
.lead-summary, .job-title {
  overflow-wrap: break-word;
  word-break: normal;
}
.entry-title, .job-title-main, .degree-name {
  flex: 1 1 auto;
  min-width: 0;
}
.entry-date, .date-range {
  flex-shrink: 0;
}
""".strip()


def _render_resume_html(template_id: str, data: dict) -> str:
    template = _jinja_env().get_template(f"{template_id}/index.html")
    rendered = template.render(**data)
    rendered = re.sub(
        r'<link[^>]+rel=["\']stylesheet["\'][^>]*>\s*',
        "",
        rendered,
        flags=re.IGNORECASE,
    )
    css_path = _backend_templates_root() / template_id / "style.css"
    css = _RESUME_RENDER_BASE_CSS + "\n" + css_path.read_text(encoding="utf-8")
    style_tag = f"<style>\n{css}\n</style>\n"
    if "</head>" in rendered:
        return rendered.replace("</head>", style_tag + "</head>", 1)
    return style_tag + rendered


def _html_to_pdf_bytes(html: str) -> bytes:
    weasy_err = None
    try:
        from weasyprint import HTML

        return HTML(string=html).write_pdf()
    except Exception as exc:
        weasy_err = f"weasyprint failed: {exc}"

    pdfkit_err = None
    try:
        import pdfkit

        return pdfkit.from_string(
            html,
            False,
            options={
                "enable-local-file-access": None,
                "quiet": None,
            },
        )
    except Exception as exc:
        pdfkit_err = f"pdfkit failed: {exc}"

    raise RuntimeError(
        "PDF generation is not available. "
        f"{weasy_err} {pdfkit_err} "
        "Install WeasyPrint with system dependencies, or install pdfkit + wkhtmltopdf."
    )


def _gemini_generate_text(prompt: str) -> str:
    # Try Gemini / Google Generative Language first
    api_key = get_setting("NEXTHIRE_AI_API_KEY") or get_setting("GOOGLE_API_KEY")
    model = get_setting("NEXTHIRE_AI_MODEL", "gemini-1.5-flash")

    def _call_gemini():
        if not api_key:
            raise RuntimeError("Gemini API key not configured. Set NEXTHIRE_AI_API_KEY or GOOGLE_API_KEY.")
        endpoint = (
            f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
            f"?key={api_key}"
        )
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.4},
        }

        req = urllib.request.Request(
            endpoint,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=40) as resp:
                raw = resp.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace") if exc.fp else str(exc)
            raise RuntimeError(f"Gemini AI request failed: {raw}")

        data = json.loads(raw or "{}")
        candidates = data.get("candidates") or []
        if not candidates:
            raise RuntimeError("Gemini AI returned no candidates.")

        parts = (((candidates[0] or {}).get("content") or {}).get("parts") or [])
        text = "".join((p.get("text") or "") for p in parts if isinstance(p, dict)).strip()
        if not text:
            raise RuntimeError("Gemini AI returned empty text.")
        return text

    def _groq_generate_text(prompt: str) -> str:
        groq_key = get_setting("GROQ_API_KEY")
        if not groq_key:
            raise RuntimeError("Groq API key not configured. Set GROQ_API_KEY.")

        groq_model = get_setting("GROQ_MODEL", "llama-3.3-70b-versatile")
        endpoint = "https://api.groq.com/openai/v1/chat/completions"
        payload = {
            "model": groq_model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.4,
        }

        req = urllib.request.Request(
            endpoint,
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {groq_key}",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=40) as resp:
                raw = resp.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace") if exc.fp else str(exc)
            raise RuntimeError(f"Groq AI request failed: {raw}")

        data = json.loads(raw or "{}")
        choices = data.get("choices") or []
        if choices:
            first = choices[0]
            # Chat-completions style
            msg = first.get("message") or first.get("delta") or {}
            content = ""
            if isinstance(msg, dict):
                # Some providers nest content differently
                content = msg.get("content") or msg.get("text") or ""
            if not content:
                # Try common fallbacks
                content = first.get("text") or first.get("message", {}).get("content") or ""
            text = (content or "").strip()
            if text:
                return text

        # Fallback: try top-level 'text' or 'choices[0].text'
        text_fallback = (data.get("text") or (choices[0].get("text") if choices else None) or "").strip()
        if text_fallback:
            return text_fallback

        raise RuntimeError("Groq AI returned no usable text.")

    # First try Gemini; on failure, attempt Groq if configured
    try:
        return _call_gemini()
    except Exception as gem_exc:
        logger.warning(f"Gemini call failed: {gem_exc}. Attempting Groq fallback.")
        try:
            return _groq_generate_text(prompt)
        except Exception as groq_exc:
            logger.error(f"Groq fallback failed: {groq_exc}")
            # Raise combined error for visibility
            raise RuntimeError(f"AI request failed. Gemini error: {gem_exc}; Groq error: {groq_exc}")


def _coerce_list(value) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _safe_resume_payload(payload: dict) -> dict:
    return {
        "name": payload.get("name", "") or "",
        "email": payload.get("email", "") or "",
        "phone": payload.get("phone", "") or "",
        "linkedin": payload.get("linkedin", "") or "",
        "objective": payload.get("objective", "") or "",
        "skills": _coerce_list(payload.get("skills")),
        "experience": _coerce_list(payload.get("experience")),
        "education": _coerce_list(payload.get("education")),
    }


def _truncate_text(value, limit: int = 700) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "..."


def _strip_code_fences(text: str) -> str:
    cleaned = str(text or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _extract_json_object(text: str) -> dict:
    cleaned = _strip_code_fences(text)
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("AI response did not contain a JSON object.")
    return json.loads(cleaned[start : end + 1])


def _safe_int(value, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        try:
            return int(float(value))
        except Exception:
            return default


def _average(values: list[float]) -> float | None:
    if not values:
        return None
    return sum(values) / len(values)


def _normalize_interview_history(value) -> list[dict]:
    history = _coerce_list(value)
    normalized: list[dict] = []
    for item in history:
        if not isinstance(item, dict):
            continue
        strengths = [str(x).strip() for x in _coerce_list(item.get("strengths")) if str(x).strip()]
        weaknesses = [str(x).strip() for x in _coerce_list(item.get("weaknesses")) if str(x).strip()]
        suggestions = [str(x).strip() for x in _coerce_list(item.get("suggestions")) if str(x).strip()]
        normalized.append(
            {
                "question": _truncate_text(item.get("question") or item.get("current_question") or "", 700),
                "answer": _truncate_text(item.get("answer") or item.get("user_answer") or "", 1000),
                "score": item.get("score"),
                "strengths": strengths,
                "weaknesses": weaknesses,
                "suggestions": suggestions,
                "ideal_answer": _truncate_text(item.get("ideal_answer") or item.get("idealAnswer") or "", 1000),
            }
        )
    return normalized


def _interview_settings(payload: dict) -> dict:
    role = str(payload.get("role") or payload.get("job_role") or payload.get("position") or "").strip()
    experience = str(payload.get("experience_level") or payload.get("experience") or payload.get("level") or "").strip()
    interview_type = str(payload.get("interview_type") or payload.get("type") or payload.get("mode") or "").strip()
    question_count = _safe_int(payload.get("question_count") or payload.get("questions") or 10, 10)
    question_number = _safe_int(payload.get("question_number") or len(_coerce_list(payload.get("interview_history"))) + 1, 1)
    return {
        "role": role or "Generalist Candidate",
        "experience_level": experience or "Intermediate",
        "interview_type": interview_type or "Mixed",
        "question_count": max(1, question_count),
        "question_number": max(1, question_number),
    }


def _history_summary(history: list[dict]) -> tuple[str, str]:
    if not history:
        return "No previous interview history.", "balanced"

    scores = [float(item["score"]) for item in history if isinstance(item.get("score"), (int, float))]
    avg_score = _average(scores)
    if avg_score is None:
        trend = "balanced"
        guidance = "There are prior questions, but no scored feedback yet. Keep the next question realistic and moderately challenging."
    elif avg_score < 4:
        trend = "easier"
        guidance = "The candidate is struggling. Ask a simpler follow-up question that builds confidence and reveals fundamentals."
    elif avg_score < 7:
        trend = "balanced"
        guidance = "The candidate is doing okay. Keep the question practical, specific, and role-relevant."
    else:
        trend = "harder"
        guidance = "The candidate is performing strongly. Increase difficulty and ask for trade-offs, depth, or concrete examples."

    latest = history[-1]
    if latest.get("weaknesses"):
        guidance += f" Recent weakness to address: {latest['weaknesses'][0]}."
    elif latest.get("strengths"):
        guidance += f" Recent strength to build on: {latest['strengths'][0]}."

    return guidance, trend


def _build_history_prompt(history: list[dict]) -> str:
    if not history:
        return "- No previous questions yet."

    lines = []
    for index, item in enumerate(history[-8:], start=max(1, len(history) - 7)):
        strengths = ", ".join(item.get("strengths") or []) or "None"
        weaknesses = ", ".join(item.get("weaknesses") or []) or "None"
        suggestions = ", ".join(item.get("suggestions") or []) or "None"
        lines.append(
            f"{index}. Q: {item.get('question') or ''}\n"
            f"   A: {item.get('answer') or ''}\n"
            f"   Score: {item.get('score') if item.get('score') is not None else 'N/A'}\n"
            f"   Strengths: {strengths}\n"
            f"   Weaknesses: {weaknesses}\n"
            f"   Suggestions: {suggestions}"
        )
    return "\n".join(lines)


def _normalize_question_signature(text: str) -> str:
    cleaned = str(text or "").strip().lower()
    cleaned = re.sub(rf"[{re.escape(string.punctuation)}]", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned


def _collect_previous_questions(payload: dict, history: list[dict]) -> list[str]:
    candidates: list[str] = []
    for key in ("askedQuestions", "asked_questions", "previous_questions", "question_history"):
        value = payload.get(key)
        if isinstance(value, list):
            candidates.extend(str(item).strip() for item in value if str(item).strip())
        elif isinstance(value, str) and value.strip():
            candidates.append(value.strip())

    for item in history:
        question = str(item.get("question") or item.get("current_question") or "").strip()
        if question:
            candidates.append(question)

    seen = set()
    unique_questions: list[str] = []
    for question in candidates:
        signature = _normalize_question_signature(question)
        if not signature or signature in seen:
            continue
        seen.add(signature)
        unique_questions.append(question)
    return unique_questions


def _is_duplicate_question(candidate: str, previous_questions: list[str]) -> bool:
    candidate_signature = _normalize_question_signature(candidate)
    if not candidate_signature:
        return True

    for previous in previous_questions:
        previous_signature = _normalize_question_signature(previous)
        if not previous_signature:
            continue
        if previous_signature == candidate_signature:
            return True
        if len(candidate_signature) > 20 and len(previous_signature) > 20:
            if difflib.SequenceMatcher(None, candidate_signature, previous_signature).ratio() >= 0.88:
                return True
    return False


def _fallback_question(settings: dict, trend: str, previous_questions: list[str] | None = None) -> dict:
    role = settings["role"]
    experience = settings["experience_level"]
    interview_type = settings["interview_type"]
    question_number = settings.get("question_number", 1)
    previous_questions = previous_questions or []
    previous_signatures = {_normalize_question_signature(q) for q in previous_questions}

    # Vary the fallback questions to avoid repetition
    fallback_options = {
        "easier": [
            f"Can you explain what {role}s typically need to understand about {interview_type.lower()} scenarios?",
            f"Walk me through a day-to-day task you'd handle in a {role} role.",
            f"What's a fundamental concept every {role} should master?",
            f"How would you describe the main responsibility of a {role} to someone new to the field?",
        ],
        "harder": [
            f"Describe a complex technical challenge you'd face as a {role} and how you'd approach it.",
            f"Tell me about a difficult trade-off decision between quality, speed, and cost in a {role} context.",
            f"How would you design a solution that scales for a large {role} challenge?",
            f"What edge cases or failure modes would you plan for when solving a {role}-level problem?",
        ],
        "maintain": [
            f"What's a recent project or task where you applied {interview_type.lower()} thinking as a {role}?",
            f"How do you balance technical depth with practical solutions in a {role} role?",
            f"Tell me about a time when your {role} skills made a measurable difference.",
            f"What tools or methodologies do you find most effective for {role} work?",
            f"How would you handle a situation where your first approach for a {role} task did not work?",
        ],
    }

    # Select a question based on trend and question number for variety
    options = fallback_options.get(trend, fallback_options["maintain"])
    question = None
    for option in options:
        if _normalize_question_signature(option) not in previous_signatures:
            question = option
            break
    if not question:
        question_index = (question_number - 1) % len(options)
        question = options[question_index]

    if trend == "easier":
        difficulty = "Beginner"
    elif trend == "harder":
        difficulty = "Advanced"
    else:
        difficulty = experience

    return {
        "question": question,
        "difficulty": difficulty,
        "focus_area": interview_type,
        "follow_up_style": trend,
        "trend": trend,
        "question_number": settings["question_number"],
        "total_questions": settings["question_count"],
    }


def _analyze_question_type(question: str) -> dict:
    """Analyze question content to determine type and focus areas."""
    q_lower = question.lower()
    
    analysis = {
        "type": "general",
        "is_technical": False,
        "is_behavioral": False,
        "is_communication": False,
        "is_architecture": False,
        "is_problem_solving": False,
        "is_process": False,
        "mentions_tools": [],
    }
    
    # Question type detection
    if any(keyword in q_lower for keyword in ["tell me", "describe", "walk me through", "explain your", "how would you", "what's a time"]):
        analysis["is_behavioral"] = True
        analysis["type"] = "behavioral"
    
    if any(keyword in q_lower for keyword in ["design", "architecture", "system", "scale", "build", "implement", "technology stack"]):
        analysis["is_architecture"] = True
        if analysis["type"] == "general":
            analysis["type"] = "technical"
    
    if any(keyword in q_lower for keyword in ["challenge", "problem", "trade-off", "edge case", "handle", "approach", "solve", "debug"]):
        analysis["is_problem_solving"] = True
        if analysis["type"] == "general":
            analysis["type"] = "technical"
    
    if any(keyword in q_lower for keyword in ["process", "method", "approach", "workflow", "communicate", "present", "explain"]):
        analysis["is_communication"] = True
    
    if any(keyword in q_lower for keyword in ["technical", "code", "api", "database", "algorithm", "performance", "optimization"]):
        analysis["is_technical"] = True
        if analysis["type"] == "general":
            analysis["type"] = "technical"
    
    if any(keyword in q_lower for keyword in ["team", "collaborate", "conflict", "disagreement", "feedback", "deadline", "pressure"]):
        analysis["is_process"] = True
    
    return analysis


def _extract_answer_concepts(answer: str, question: str) -> dict:
    """Extract technologies, concepts, and structured elements from answer."""
    answer_lower = answer.lower()
    
    concepts = {
        "technologies": [],
        "concepts": [],
        "has_examples": False,
        "has_metrics": False,
        "has_outcome": False,
        "response_length": len(answer.split()),
        "uses_star_structure": False,
    }
    
    # Common technology keywords
    tech_keywords = {
        "python", "javascript", "typescript", "java", "react", "node", "django", "flask",
        "fastapi", "sql", "mongodb", "redis", "docker", "kubernetes", "aws", "azure", "gcp",
        "git", "rest", "api", "microservices", "html", "css", "vue", "angular", "nextjs",
        "graphql", "postgresql", "mysql", "nosql", "redis", "elasticsearch", "kafka",
        "ci/cd", "jenkins", "github", "gitlab", "testing", "jest", "pytest", "junit",
        "agile", "scrum", "devops", "linux", "windows", "ssl", "jwt", "oauth", "authentication",
        "performance", "optimization", "caching", "cdn", "cdn", "load", "balance",
    }
    
    for tech in tech_keywords:
        if tech in answer_lower:
            concepts["technologies"].append(tech)
    
    # Check for structured communication patterns
    if any(indicator in answer_lower for indicator in ["situation", "task", "action", "result"]):
        concepts["uses_star_structure"] = True
    
    # Check for concrete elements
    concepts["has_examples"] = any(indicator in answer_lower for indicator in [
        "example", "for instance", "like", "such as", "specifically", "particular", "case"
    ])
    
    concepts["has_metrics"] = bool(re.search(r'\d+\s*%|\d+\s*x|\d+\s*sec|seconds|ms|milliseconds|reduce|increase|improve', answer_lower))
    
    concepts["has_outcome"] = any(indicator in answer_lower for indicator in [
        "result", "outcome", "impact", "improve", "achieve", "success", "learned", "discovered", "found"
    ])
    
    # General concept detection
    concept_keywords = {
        "scalability", "performance", "security", "maintainability", "testing", "documentation",
        "debugging", "architecture", "design pattern", "oop", "functional", "asynchronous", 
        "synchronous", "threading", "concurrency", "caching", "optimization", "refactoring",
        "deployment", "monitoring", "logging", "error handling", "validation", "sanitization",
    }
    
    for concept in concept_keywords:
        if concept in answer_lower:
            concepts["concepts"].append(concept)
    
    return concepts


def _fallback_evaluation(answer_text: str, settings: dict, question: str = "") -> dict:
    """Enhanced fallback evaluation with contextual analysis."""
    word_count = len([w for w in str(answer_text or "").split() if w.strip()])
    
    # Analyze question type for context-aware feedback
    q_analysis = _analyze_question_type(question) if question else {}
    answer_analysis = _extract_answer_concepts(answer_text, question)
    
    # Base scoring on content quality, not just length
    base_score = 5
    if word_count >= 150:
        base_score = 7
    elif word_count >= 80:
        base_score = 6
    elif word_count >= 35:
        base_score = 5
    else:
        base_score = 3
    
    # Bonus points for specific content elements
    if answer_analysis["has_metrics"]:
        base_score += 1
    if answer_analysis["has_examples"]:
        base_score += 0.5
    if answer_analysis["uses_star_structure"] and q_analysis.get("is_behavioral"):
        base_score += 1
    
    # Penalties for negative indicators
    if "i don't know" in answer_text.lower() or "not sure" in answer_text.lower():
        base_score = max(1, base_score - 2)
    if "um" in answer_text.lower() or "uh" in answer_text.lower():
        base_score = max(1, base_score - 0.5)
    
    score = max(1, min(10, int(base_score)))
    
    # Context-aware feedback generation
    strengths = []
    weaknesses = []
    suggestions = []
    
    if score >= 7:
        if answer_analysis["has_metrics"]:
            strengths.append("Good use of specific metrics and quantifiable results.")
        if answer_analysis["has_examples"]:
            strengths.append("Provided concrete examples to support your point.")
        if answer_analysis["response_length"] >= 100:
            strengths.append("Detailed and well-structured response.")
        if not strengths:
            strengths.append("Clear communication and solid technical understanding.")
    elif score >= 5:
        strengths.append("Attempted to provide a structured response.")
        if answer_analysis["has_examples"]:
            strengths.append("Included some relevant examples.")
    else:
        strengths.append("You attempted to answer the question.")
    
    # Role-specific weakness detection
    if q_analysis.get("is_technical"):
        if not answer_analysis["technologies"]:
            weaknesses.append("Missing discussion of specific technologies or tools.")
        if answer_analysis["response_length"] < 50:
            weaknesses.append("Technical explanation lacks sufficient depth.")
        if not answer_analysis["has_metrics"] and "performance" in question.lower():
            weaknesses.append("Missing measurable impact or performance metrics.")
    
    if q_analysis.get("is_behavioral"):
        if not answer_analysis["uses_star_structure"]:
            weaknesses.append("Response could benefit from STAR structure (Situation, Task, Action, Result).")
        if not answer_analysis["has_outcome"]:
            weaknesses.append("Missing the outcome or lesson learned from the situation.")
    
    if not answer_analysis["has_examples"]:
        weaknesses.append("More specific examples would strengthen your answer.")
    
    if not weaknesses:
        weaknesses.append("Add more depth to the technical or behavioral aspects.")
    
    # Context-aware suggestions
    if q_analysis.get("is_technical"):
        suggestions.append("Mention specific technologies or frameworks you'd use.")
        suggestions.append("Discuss trade-offs between different approaches.")
    elif q_analysis.get("is_behavioral"):
        suggestions.append("Use the STAR method: Situation, Task, Action, Result.")
        suggestions.append("Explain what you learned from the experience.")
    
    if not answer_analysis["has_metrics"]:
        suggestions.append("Include measurable outcomes or metrics when possible.")
    
    if answer_analysis["response_length"] < 60:
        suggestions.append("Provide more detailed context to fully address the question.")
    
    suggestions = suggestions[:3]  # Limit to top 3
    if not suggestions:
        suggestions = ["Give concrete examples.", "Explain the measurable impact."]
    
    return {
        "score": score,
        "strengths": strengths or ["You attempted a structured response."],
        "weaknesses": weaknesses or ["Add more specific detail."],
        "suggestions": suggestions,
        "ideal_answer": "A strong answer would include specific examples, measurable outcomes, and demonstrate deep understanding of the role.",
        "communication_rating": max(1, min(10, score + (1 if answer_analysis["response_length"] >= 80 else 0))),
        "technical_rating": max(1, min(10, score if q_analysis.get("is_technical") else score - 1)),
        "confidence_rating": max(1, min(10, score - 1 if score > 2 else 1)),
        "performance_level": "Excellent" if score >= 9 else "Strong" if score >= 8 else "Solid" if score >= 6 else "Developing" if score >= 4 else "Needs Support",
        "follow_up_direction": "harder" if score >= 8 else "maintain" if score >= 5 else "easier",
        "short_feedback": "Your answer has potential, but add more specific examples and measurable outcomes." if score < 7 else "Good response with solid structure and relevant details.",
    }


def _generate_interview_question_from_gemini(payload: dict) -> dict:
    settings = _interview_settings(payload)
    history = _normalize_interview_history(
        payload.get("interview_history")
        or payload.get("previous_interview_history")
        or payload.get("history")
    )
    guidance, trend = _history_summary(history)
    history_prompt = _build_history_prompt(history)
    previous_questions = _collect_previous_questions(payload, history)
    previous_questions_text = "\n".join([f"- {q}" for q in previous_questions]) if previous_questions else "(No previous questions yet)"
    covered_concepts = payload.get("covered_concepts") or []
    if isinstance(covered_concepts, list):
        covered_concepts = [str(c).strip() for c in covered_concepts if str(c).strip()]
    covered_concepts_text = ", ".join(covered_concepts) if covered_concepts else "(No concepts covered yet)"

    prompt = (
        "You are a realistic, professional recruiter/interviewer conducting a continuous interview. "
        "Your job is to ask ONE follow-up question that advances the interview naturally.\n"
        "Return valid JSON only with keys: question, difficulty, focus_area, follow_up_style. "
        "Do not include markdown, numbering, or extra commentary.\n\n"
        
        f"Role: {settings['role']}\n"
        f"Experience level: {settings['experience_level']}\n"
        f"Interview type: {settings['interview_type']}\n"
        f"Question number: {settings['question_number']} of {settings['question_count']}\n"
        f"Adaptive guidance: {guidance}\n"
        f"Suggested trend: {trend}\n\n"
        f"Interview history (most recent last):\n{history_prompt}\n\n"
        f"Previous Q&A history is mandatory context for this request.\n\n"
        
        f"PREVIOUS QUESTIONS ASKED (DO NOT REPEAT OR USE SIMILAR PHRASING):\n{previous_questions_text}\n\n"
        
        f"CONCEPTS ALREADY COVERED (MUST AVOID THESE UNLESS GOING DEEPER):\n{covered_concepts_text}\n\n"
        
        "CRITICAL ANTI-REPETITION RULES:\n"
        "- NEVER ask the exact same question or paraphrase it\n"
        "- NEVER ask about a concept that was already covered (from the list above)\n"
        "- Each question MUST explore a NEW concept, skill, or angle\n"
        "- If you must revisit a concept, make it a DEEP follow-up (e.g., 'How would you handle edge cases?' after discussing basics)\n"
        "- The interview should feel like a natural conversation that evolves and builds\n\n"
        
        "FOLLOW-UP QUESTION BEHAVIOR:\n"
        "- If the candidate performed strongly (score >= 7): Increase depth and ask for trade-offs, edge cases, or architecture decisions\n"
        "- If the candidate did okay (score 5-6): Maintain difficulty and ask about related but different concepts\n"
        "- If the candidate struggled (score < 5): Ask a simpler question about fundamentals or a different angle\n\n"
        
        "PERSONALIZATION:\n"
        "- Reference concrete items from the resume (technologies, projects, achievements) when possible\n"
        "- Build on what the candidate mentioned: 'You mentioned JWT... how would you handle token expiration?'\n"
        "- Vary question types: technical depth, architecture, problem-solving, process, trade-offs\n\n"
        
        "OUTPUT INSTRUCTIONS:\n"
        "- Return ONLY JSON with requested keys (question, difficulty, focus_area, follow_up_style)\n"
        "- The 'question' field must be a single focused question\n"
        "- Do not include any explanation or markdown in the JSON"
    )

    try:
        raw = _gemini_generate_text(prompt)
        parsed = _extract_json_object(raw)
        question = _truncate_text(parsed.get("question") or "", 1000)
        if not question:
            raise ValueError("Missing question text.")
        if _is_duplicate_question(question, previous_questions):
            raise ValueError("Gemini generated a repeated question.")
        logger.info(f"Interview question generated successfully: question_number={settings['question_number']}")
        return {
            "question": question,
            "difficulty": str(parsed.get("difficulty") or settings["experience_level"]),
            "focus_area": str(parsed.get("focus_area") or settings["interview_type"]),
            "follow_up_style": str(parsed.get("follow_up_style") or trend),
            "trend": trend,
            "question_number": settings["question_number"],
            "total_questions": settings["question_count"],
        }
    except Exception as exc:
        logger.warning(f"Interview question Gemini failed (using fallback): {exc}")
        logger.debug(f"Covered concepts: {covered_concepts}")
        logger.debug(f"Previous questions count: {len(previous_questions)}")
        return _fallback_question(settings, trend, previous_questions)


def _evaluate_interview_answer_with_gemini(payload: dict) -> dict:
    settings = _interview_settings(payload)
    history = _normalize_interview_history(
        payload.get("interview_history")
        or payload.get("previous_interview_history")
        or payload.get("history")
    )
    question = _truncate_text(payload.get("current_question") or payload.get("question") or "", 1000)
    answer = _truncate_text(payload.get("user_answer") or payload.get("answer") or "", 2000)
    history_prompt = _build_history_prompt(history)
    previous_questions = _collect_previous_questions(payload, history)
    previous_questions_text = "\n".join([f"- {q}" for q in previous_questions]) if previous_questions else "(No previous questions yet)"
    
    covered_concepts = payload.get("covered_concepts") or []
    if isinstance(covered_concepts, list):
        covered_concepts = [str(c).strip() for c in covered_concepts if str(c).strip()]
    covered_concepts_text = ", ".join(covered_concepts) if covered_concepts else "(No concepts covered yet)"

    # Analyze question and answer for context-aware evaluation
    question_analysis = _analyze_question_type(question)
    answer_analysis = _extract_answer_concepts(answer, question)

    prompt = (
        "You are a fair, rigorous, and INTELLIGENT interview evaluator acting as a senior hiring manager. "
        "Your evaluations must be ANSWER-AWARE, ROLE-AWARE, and CONTEXT-AWARE. "
        "NEVER give generic feedback like 'be more specific' or 'use examples' to everyone - customize each evaluation based on what was actually said.\n\n"
        "Return valid JSON only with keys: score, strengths, weaknesses, suggestions, ideal_answer, communication_rating, technical_rating, confidence_rating, performance_level, follow_up_direction, short_feedback.\n"
        "score must be an integer from 1 to 10. Ratings must be integers 1-10. Arrays must contain specific, actionable strings.\n"
        "Do not include markdown or extra commentary.\n\n"
        
        f"ROLE: {settings['role']}\n"
        f"EXPERIENCE LEVEL: {settings['experience_level']}\n"
        f"INTERVIEW TYPE: {settings['interview_type']}\n"
        f"QUESTION #{settings['question_number']} OF {settings['question_count']}\n\n"
        
        f"QUESTION ASKED:\n{question}\n\n"
        f"CANDIDATE'S ANSWER:\n{answer}\n\n"
        f"Interview history:\n{history_prompt}\n\n"
        f"Previous Q&A history is mandatory context for this request.\n\n"
        
        f"Previously covered concepts: {covered_concepts_text}\n\n"
        
        "INTELLIGENT EVALUATION FRAMEWORK:\n"
        "=== TECHNICAL ASSESSMENT ===\n"
        "- Is the answer technically correct? If wrong, identify the error.\n"
        "- Are specific technologies, frameworks, or concepts mentioned?\n"
        "- Does the answer show depth or is it surface-level?\n"
        "- Are there trade-offs or architectural considerations discussed?\n\n"
        
        "=== BEHAVIORAL & COMMUNICATION ASSESSMENT ===\n"
        "- Is the answer structured and easy to follow?\n"
        "- Does it use STAR structure (Situation-Task-Action-Result) when relevant?\n"
        "- Are there concrete examples and specific metrics/outcomes?\n"
        "- Is the language professional and confident?\n\n"
        
        "=== ROLE-SPECIFIC CRITERIA ===\n"
        "FOR BACKEND/FRONTEND/FULLSTACK:\n"
        "- Did they discuss scalability, performance, or optimization?\n"
        "- Did they mention testing, debugging, or maintenance?\n"
        "- Did they show awareness of security, validation, or error handling?\n\n"
        
        "FOR HR/BEHAVIORAL:\n"
        "- Does the answer demonstrate collaboration and communication?\n"
        "- Are soft skills (conflict resolution, leadership, adaptation) evident?\n"
        "- Is there evidence of learning from the experience?\n\n"
        
        "FOR DATA/ANALYTICS:\n"
        "- Were data sources, methodologies, or tools discussed?\n"
        "- Did they mention metrics, KPIs, or insights?\n"
        "- Was there discussion of data validation or accuracy?\n\n"
        
        "=== SCORE LOGIC ===\n"
        "10: Expert-level depth, all key aspects covered, concrete examples, leadership demonstration\n"
        "8-9: Strong understanding, specific examples, good structure, minor gaps\n"
        "6-7: Solid basic answer, some examples, generally on track, could go deeper\n"
        "4-5: Surface-level, vague, missing key concepts or specific examples\n"
        "1-3: Incorrect, evasive, or insufficient response\n\n"
        
        "=== STRENGTH DETECTION ===\n"
        "Find SPECIFIC strengths in THEIR answer:\n"
        "- 'You mentioned X technology, which shows knowledge of Y'\n"
        "- 'Your example demonstrates understanding of Z pattern'\n"
        "- 'You correctly identified trade-off between A and B'\n"
        "- 'STAR structure was clear: you explained situation, action, and result'\n"
        "- 'You quantified impact: reduced latency by X%' or similar metrics\n\n"
        
        "=== WEAKNESS DETECTION ===\n"
        "Find SPECIFIC weaknesses based on role and question:\n"
        "- 'You didn't discuss scalability, which is crucial for Backend roles'\n"
        "- 'Missing mention of error handling/edge cases'\n"
        "- 'The outcome/result wasn't explained'\n"
        "- 'Vague on specific technologies that Backend Developers should know'\n"
        "- 'Didn't address the actual question directly'\n\n"
        
        "=== SUGGESTIONS MUST BE SPECIFIC ===\n"
        "- Reference what they said or didn't say\n"
        "- 'You mentioned React but didn't discuss state management - consider exploring Redux, Context API, or Zustand'\n"
        "- 'When discussing API design, include error handling and authentication patterns'\n"
        "- NOT: 'Use concrete examples' (they may have already done this)\n"
        "- Suggest NEW concepts not yet covered, prioritize concepts from the covered list to avoid\n\n"
        
        "=== SHORT_FEEDBACK MUST BE SPECIFIC ===\n"
        "- If score >= 8: 'Strong answer on [topic]. Your mention of [specific detail] shows solid understanding.'\n"
        "- If score 5-7: 'Good foundation. Adding [specific concept] would strengthen your answer.'\n"
        "- If score < 5: 'Consider focusing on [specific weakness]. A recruiter would expect [role-specific detail].'\n"
        "- NEVER generic like 'needs more specificity' (FORBIDDEN - customize to their answer)\n\n"
        
        "=== FOLLOW-UP DIRECTION ===\n"
        "- If score >= 8: 'harder' (ask edge cases, trade-offs, architecture)\n"
        "- If score 5-7: 'maintain' (similar difficulty, different angle or concept)\n"
        "- If score < 5: 'easier' (fundamentals, related basics)\n\n"
        
        "OUTPUT: Return ONLY JSON with ALL requested keys."
    )

    try:
        raw = _gemini_generate_text(prompt)
        parsed = _extract_json_object(raw)
        score = _safe_int(parsed.get("score"), 5)
        strengths = [str(x).strip() for x in _coerce_list(parsed.get("strengths")) if str(x).strip()]
        weaknesses = [str(x).strip() for x in _coerce_list(parsed.get("weaknesses")) if str(x).strip()]
        suggestions = [str(x).strip() for x in _coerce_list(parsed.get("suggestions")) if str(x).strip()]
        ideal_answer = _truncate_text(parsed.get("ideal_answer") or parsed.get("idealAnswer") or "", 1400)
        communication_rating = _safe_int(parsed.get("communication_rating"), score)
        technical_rating = _safe_int(parsed.get("technical_rating"), score)
        confidence_rating = _safe_int(parsed.get("confidence_rating"), score)
        performance_level = str(parsed.get("performance_level") or "Solid")
        follow_up_direction = str(parsed.get("follow_up_direction") or "maintain")
        short_feedback = _truncate_text(parsed.get("short_feedback") or "", 500)
        if not short_feedback:
            short_feedback = "Your answer shows understanding. Consider adding more specific examples or technical depth."

        follow_up_direction = follow_up_direction if follow_up_direction in {"easier", "maintain", "harder"} else "maintain"

        return {
            "score": max(1, min(10, score)),
            "strengths": strengths or ["You provided a structured response."],
            "weaknesses": weaknesses or ["Add more specific technical or behavioral detail."],
            "suggestions": suggestions or ["Elaborate on your approach.", "Include measurable outcomes."],
            "ideal_answer": ideal_answer or "A strong answer would combine specific examples, technical depth, and measurable impact.",
            "communication_rating": max(1, min(10, communication_rating)),
            "technical_rating": max(1, min(10, technical_rating)),
            "confidence_rating": max(1, min(10, confidence_rating)),
            "performance_level": performance_level,
            "follow_up_direction": follow_up_direction,
            "short_feedback": short_feedback,
            "question": question,
        }
    except Exception as exc:
        logger.warning(f"Interview evaluation fallback: {exc}")
        fallback = _fallback_evaluation(answer, settings, question)
        fallback["question"] = question
        return fallback


@app.route("/api/render-resume-preview", methods=["POST"])
def render_resume_preview():
    try:
        payload = request.get_json(silent=True) or {}
        template_id = (payload.get("template_id") or payload.get("template") or "").strip()
        data = payload.get("data") if isinstance(payload.get("data"), dict) else payload

        if not template_id:
            return _error("Missing template_id.", 400)
        if not _template_exists(template_id):
            return _error(f"Template not found: {template_id}", 404)

        resume_data = _safe_resume_payload(data if isinstance(data, dict) else {})
        html = _render_resume_html(template_id, resume_data)
        return app.response_class(html, mimetype="text/html; charset=utf-8")
    except Exception as exc:
        logger.error(f"Resume preview error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/upload-resume", methods=["POST"])
def upload_resume():
    try:
        logger.info("Resume upload request received")
        if 'resume' not in request.files:
            return _error('No file part named "resume" provided.', 400)

        resume_file = request.files.get('resume')
        if not resume_file or not resume_file.filename:
            return _error('No resume file uploaded.', 400)

        filename = resume_file.filename or 'resume'
        content = resume_file.read()
        if not content:
            return _error('Uploaded file is empty.', 400)

        lower = filename.lower()
        parsed_text = ''
        if lower.endswith('.pdf'):
            parsed_text = _extract_text_from_pdf(content)
            source = 'pdf'
        elif lower.endswith('.docx'):
            parsed_text = _extract_text_from_docx(content)
            source = 'docx'
        else:
            return _error('Unsupported file type. Accepted: PDF, DOCX', 400)

        # Limit length returned to frontend to a reasonable size
        preview = (parsed_text or '')[:10000]
        logger.info(f"Parsed resume ({source}): {len(preview)} chars")
        return jsonify({"filename": filename, "parsed_text": preview, "source": source})
    except Exception as exc:
        logger.error(f"Resume upload error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/generate-question", methods=["POST"])
@app.route("/generate-question", methods=["POST"])
def generate_question():
    try:
        logger.info("Interview question request received")
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            payload = {}
        return jsonify(_generate_interview_question_from_gemini(payload))
    except Exception as exc:
        logger.error(f"Interview question error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/evaluate-answer", methods=["POST"])
@app.route("/evaluate-answer", methods=["POST"])
def evaluate_answer():
    try:
        logger.info("Interview evaluation request received")
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            payload = {}
        return jsonify(_evaluate_interview_answer_with_gemini(payload))
    except Exception as exc:
        logger.error(f"Interview evaluation error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/interview-summary", methods=["POST"])
@app.route("/interview-summary", methods=["POST"])
def interview_summary():
    try:
        logger.info("Interview summary request received")
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            payload = {}

        history = _normalize_interview_history(
            payload.get("interview_history")
            or payload.get("previous_interview_history")
            or payload.get("history")
        )
        settings = _interview_settings(payload)
        score_values = [float(item.get("score")) for item in history if isinstance(item.get("score"), (int, float))]
        communication_values = [float(item.get("communication_rating")) for item in history if isinstance(item.get("communication_rating"), (int, float))]
        technical_values = [float(item.get("technical_rating")) for item in history if isinstance(item.get("technical_rating"), (int, float))]
        confidence_values = [float(item.get("confidence_rating")) for item in history if isinstance(item.get("confidence_rating"), (int, float))]
        avg_score = _average(score_values) or 0

        summary = {
            "overall_score": round(avg_score, 1),
            "performance_level": (
                "Not Started"
                if not history
                else "Excellent"
                if avg_score >= 8.5
                else "Strong"
                if avg_score >= 7
                else "Solid"
                if avg_score >= 5.5
                else "Developing"
                if avg_score >= 4
                else "Needs Support"
            ),
            "communication_rating": round(_average(communication_values) or 0, 1),
            "technical_rating": round(_average(technical_values) or 0, 1),
            "confidence_rating": round(_average(confidence_values) or 0, 1),
            "strong_areas": [],
            "areas_needing_improvement": [],
            "final_feedback_message": "",
            "role": settings["role"],
            "experience_level": settings["experience_level"],
            "interview_type": settings["interview_type"],
        }

        strong_areas: list[str] = []
        improvement_areas: list[str] = []
        for item in history:
            for value in _coerce_list(item.get("strengths")):
                text = _safe_str(value)
                if text and text not in strong_areas:
                    strong_areas.append(text)
            for value in _coerce_list(item.get("weaknesses")) + _coerce_list(item.get("suggestions")):
                text = _safe_str(value)
                if text and text not in improvement_areas:
                    improvement_areas.append(text)

        summary["strong_areas"] = strong_areas[:5]
        summary["areas_needing_improvement"] = improvement_areas[:6]
        summary["final_feedback_message"] = (
            _safe_str(history[-1].get("short_feedback"), "")
            if history
            else "You completed the interview practice session. Repeat the interview to improve your score and depth."
        )

        # Hiring recommendation and improvement roadmap
        rec = "Consider for next-stage interviews"
        if avg_score >= 8.5:
            rec = "Strong hire" 
        elif avg_score >= 7:
            rec = "Consider hire (next-stage)"
        elif avg_score >= 5.5:
            rec = "Potential hire with coaching"
        elif avg_score >= 4:
            rec = "Not yet ready; more practice recommended"
        else:
            rec = "Not recommended at this time"

        roadmap = []
        for area in summary["areas_needing_improvement"]:
            roadmap.append(f"Practice: {area}. Work on examples and concrete outcomes.")

        summary["hiring_recommendation"] = rec
        summary["improvement_roadmap"] = roadmap[:6]

        return jsonify(summary)
    except Exception as exc:
        logger.error(f"Interview summary error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/test", methods=["GET"])
def test():
    logger.info("Health check endpoint called")
    return jsonify({"message": "Backend working!"})


@app.route("/api/ats-score", methods=["POST"])
def ats_score():
    try:
        logger.info("ATS score request received")
        resume_text = ""
        jd_text = ""
        resume_source = "text"
        jd_source = "text"
        job_title = ""
        industry = ""

        if request.is_json:
            payload = request.get_json(silent=True) or {}
            resume_text = payload.get("resume_text", "") or ""
            jd_text = payload.get("jd_text", "") or payload.get("job_description", "") or ""
            job_title = payload.get("job_title", "") or ""
            industry = payload.get("industry", "") or ""
        else:
            form = request.form or {}
            resume_text = form.get("resume_text", "") or ""
            jd_text = form.get("jd_text", "") or ""

            resume_file = request.files.get("resume_pdf")
            jd_file = request.files.get("jd_pdf")

            if resume_file and resume_file.filename:
                resume_bytes = resume_file.read()
                if not resume_bytes:
                    logger.warning("Empty resume PDF received")
                    return _error("Resume PDF is empty.")
                resume_text = _extract_text_from_pdf(resume_bytes)
                resume_source = "pdf"

            if jd_file and jd_file.filename:
                jd_bytes = jd_file.read()
                if not jd_bytes:
                    logger.warning("Empty JD PDF received")
                    return _error("Job description PDF is empty.")
                jd_text = _extract_text_from_pdf(jd_bytes)
                jd_source = "pdf"

            if jd_text and jd_source == "text":
                jd_source = "text"

        # Accept job title and industry from form fields
        try:
            job_title = job_title or (request.form.get("job_title") or "")
            industry = industry or (request.form.get("industry") or "")
        except Exception:
            pass

        if not str(resume_text).strip():
            logger.warning("ATS score request missing resume")
            return _error("No resume provided. Send resume_text or resume_pdf.")
        if not str(jd_text).strip():
            logger.warning("ATS score request missing job description")
            return _error("No job description provided. Send jd_text / job_description or jd_pdf.")

        # PRIMARY: Use Gemini AI for ATS scoring (most accurate)
        # Fallback to local calculation only if Gemini fails
        try:
            prompt = (
                "You are an expert ATS (Applicant Tracking System) analyst. Analyze this resume against the job description "
                "using these proven methodologies:\n\n"
                
                "ANALYSIS FRAMEWORK:\n"
                "1. KEYWORD FREQUENCY & CONTEXTUAL MAPPING:\n"
                "   - Direct Skills Match: Exact matches for core requirements (HTML, CSS, JavaScript, etc.)\n"
                "   - Desired Skills Match: Plus/nice-to-have keywords (React, frameworks, etc.)\n"
                "   - Verb Alignment: Check for action verbs from JD (developing, maintaining, collaborating, etc.)\n"
                "   - Project Relevance: Does technical stack show exceeding competency for the role?\n\n"
                
                "2. ROLE ELIGIBILITY & HIERARCHY VERIFICATION:\n"
                "   - Educational Background: Does degree/major match requirements?\n"
                "   - Experience Level: Does career stage (intern, junior, mid-level) align with JD target?\n"
                "   - Graduation Timeline: Is timing appropriate (final semester, recent grad, etc.)?\n\n"
                
                "3. STRUCTURAL PARSABILITY:\n"
                "   - Section Headers: Are standard sections (Skills, Education, Experience, Projects) present?\n"
                "   - Contact Info: Email, phone, location clearly visible?\n"
                "   - Formatting: Professional layout, ATS-friendly formatting?\n\n"
                
                "4. TECHNICAL DEPTH ASSESSMENT:\n"
                "   - Technical Stack: Does it exceed minimum requirements?\n"
                "   - Project Complexity: Quality and relevance of projects listed?\n"
                "   - Skill Progression: Evidence of growth and learning?\n\n"
                
                f"JOB TITLE: {job_title}\n"
                f"INDUSTRY: {industry}\n\n"
                "JOB DESCRIPTION:\n---\n" + jd_text[:3000] + "\n---\n\n"
                "RESUME TEXT:\n---\n" + resume_text[:3000] + "\n---\n\n"
                
                "Return ONLY valid JSON with no extra text:\n"
                "{\n"
                "  \"score\": <0-100 integer - overall ATS match>,\n"
                "  \"keyword_match_score\": <0-100>,\n"
                "  \"eligibility_score\": <0-100>,\n"
                "  \"parsability_score\": <0-100>,\n"
                "  \"technical_depth_score\": <0-100>,\n"
                "  \"matched_keywords\": [<array of matched skills>],\n"
                "  \"missing_keywords\": [<array of missing but important skills>],\n"
                "  \"formatting_issues\": [<array of formatting concerns like 'two-column layout', 'images detected', 'long paragraphs', etc.>],\n"
                "  \"core_requirements_met\": <true/false>,\n"
                "  \"plus_skills_present\": [<bonus skills that are present>],\n"
                "  \"education_match\": <brief assessment>,\n"
                "  \"experience_fit\": <brief assessment>,\n"
                "  \"strengths\": [<top 3 resume strengths>],\n"
                "  \"gaps\": [<top 3 areas to improve>],\n"
                "  \"overall_assessment\": <one sentence summary>\n"
                "}\n\n"
                "SCORING CALIBRATION:\n"
                "- Score 80-90: Strong match (core + plus skills, excellent project fit)\n"
                "- Score 90+: Excellent match (all requirements met, exceeds expectations)\n"
                "- Be generous with scores for good candidates, similar to how a human recruiter or ChatGPT/Gemini web would evaluate.\n"
                "- A score of 70-80 is 'Good', 80-90 is 'Great', 90+ is 'Outstanding'."
            )

            gemini_response = _gemini_generate_text(prompt)
            logger.info(f"Gemini raw response: {gemini_response[:500]}")
            
            # Extract JSON from response
            gemini_response = gemini_response.strip()
            json_start = gemini_response.find('{')
            json_end = gemini_response.rfind('}')
            
            if json_start >= 0 and json_end > json_start:
                json_text = gemini_response[json_start:json_end+1]
                parsed = json.loads(json_text)
                
                # Build comprehensive result from Gemini
                result = {
                    "ats_score": round(float(parsed.get('score', 0)), 2),
                    "keyword_score": round(float(parsed.get('keyword_match_score', 0)), 2),
                    "semantic_score": round(float(parsed.get('technical_depth_score', 0)), 2),
                    "keyword_match_score": round(float(parsed.get('keyword_match_score', 0)), 2),
                    "eligibility_score": round(float(parsed.get('eligibility_score', 0)), 2),
                    "parsability_score": round(float(parsed.get('parsability_score', 0)), 2),
                    "technical_depth_score": round(float(parsed.get('technical_depth_score', 0)), 2),
                    "matched_keywords": [str(x).strip() for x in parsed.get('matched_keywords', []) if x],
                    "missing_keywords": [str(x).strip() for x in parsed.get('missing_keywords', []) if x],
                    "formatting_issues": [str(x).strip() for x in parsed.get('formatting_issues', []) if x],
                    "core_requirements_met": parsed.get('core_requirements_met', False),
                    "plus_skills_present": [str(x).strip() for x in parsed.get('plus_skills_present', []) if x],
                    "education_match": str(parsed.get('education_match', '')),
                    "experience_fit": str(parsed.get('experience_fit', '')),
                    "strengths": [str(x).strip() for x in parsed.get('strengths', []) if x],
                    "gaps": [str(x).strip() for x in parsed.get('gaps', []) if x],
                    "overall_assessment": str(parsed.get('overall_assessment', '')),
                    "resume_source": resume_source,
                    "jd_source": jd_source,
                    "analysis_method": "Gemini AI (Expert ATS Analysis)",
                }
                
                try:
                    result['resume_text_preview'] = (resume_text or "")[:2000]
                except Exception:
                    result['resume_text_preview'] = ""
                
                logger.info(f"ATS Score (Gemini): {result.get('ats_score')}")
                return jsonify(result)
                
        except json.JSONDecodeError as je:
            logger.warning(f"Gemini JSON parse error: {je}")
        except Exception as gemini_err:
            logger.warning(f"Gemini ATS analysis error: {gemini_err}")

        # FALLBACK: Use local calculation if Gemini fails
        logger.info("Falling back to local ATS calculation")
        local_result = calculate_advanced_ats(resume_text, jd_text, resume_source, jd_source)
        
        try:
            local_result['resume_text_preview'] = (resume_text or "")[:2000]
        except Exception:
            local_result['resume_text_preview'] = ""
        
        local_result['analysis_method'] = "Local Analysis (Fallback)"
        logger.info(f"ATS Score (Local Fallback): {local_result.get('ats_score')}")
        return jsonify(local_result)
    except Exception as exc:
        logger.error(f"ATS score error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/generate-summary", methods=["POST"])
def generate_summary():
    try:
        logger.info("Summary generation request received")
        payload = request.get_json(silent=True) or {}
        skills = payload.get("skills", [])
        experience = payload.get("experience", [])

        skills_text = ", ".join(
            [
                f"{(s or {}).get('name', '')}".strip()
                for s in skills
                if isinstance(s, dict) and (s or {}).get("name")
            ]
        )
        exp_lines = []
        for e in experience if isinstance(experience, list) else []:
            if not isinstance(e, dict):
                continue
            title = (e.get("title") or "").strip()
            company = (e.get("company") or "").strip()
            desc = (e.get("description") or "").strip()
            line = " - ".join([p for p in [title, company] if p])
            if desc:
                line = f"{line}: {desc}" if line else desc
            if line:
                exp_lines.append(line)

        prompt = (
            "You are an expert resume writer. Write a concise professional summary (2-4 lines) "
            "for a resume. Keep it ATS-friendly, quantified where possible, and avoid emojis.\n\n"
            f"Skills: {skills_text}\n"
            "Experience:\n"
            + ("\n".join(f"- {x}" for x in exp_lines) if exp_lines else "- (none provided)")
            + "\n\nOutput only the summary text."
        )

        summary = _gemini_generate_text(prompt)
        logger.info("Summary generated successfully")
        return jsonify({"summary": summary})
    except Exception as exc:
        logger.error(f"Summary generation error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/improve-text", methods=["POST"])
def improve_text():
    try:
        logger.info("Text improvement request received")
        payload = request.get_json(silent=True) or {}
        raw_text = (payload.get("text") or "").strip()
        if not raw_text:
            logger.warning("Text improvement: empty text provided")
            return _error("Missing text.", 400)

        prompt = (
            "You are an expert resume writer. Rewrite the following experience description into "
            "professional, ATS-friendly bullet points. Use action verbs, include impact/metrics when possible, "
            "and keep it concise.\n\n"
            "Return 3-6 bullets, each on a new line, starting with '- '. Output only the bullets.\n\n"
            f"TEXT:\n{raw_text}"
        )

        improved = _gemini_generate_text(prompt)
        
        lines = [ln.strip() for ln in improved.splitlines() if ln.strip()]
        bullets = []
        for ln in lines:
            cleaned = ln.lstrip("•*- ").strip()
            if cleaned:
                bullets.append(cleaned)

        logger.info(f"Text improved: {len(bullets)} bullets generated")
        return jsonify({"improved_text": improved.strip(), "bullets": bullets})
    except Exception as exc:
        logger.error(f"Text improvement error: {str(exc)}")
        return _error(str(exc), 500)


@app.route("/api/generate-resume", methods=["POST"])
def generate_resume():
    try:
        logger.info("Resume generation request received")
        payload = request.get_json(silent=True) or {}
        template_id = (payload.get("template_id") or "").strip() or (payload.get("template") or "").strip()
        data = payload.get("data") if isinstance(payload.get("data"), dict) else payload

        if not template_id:
            logger.warning("Resume generation: missing template_id")
            return _error("Missing template_id.", 400)
        if not _template_exists(template_id):
            logger.warning(f"Resume generation: template not found: {template_id}")
            return _error(f"Template not found: {template_id}", 404)

        resume_data = _safe_resume_payload(data if isinstance(data, dict) else {})

        html = _render_resume_html(template_id, resume_data)
        pdf_bytes = _html_to_pdf_bytes(html)
        
        filename = f"{(resume_data.get('name') or 'resume').strip().replace(' ', '_')}.pdf"
        logger.info(f"Resume generated successfully: {filename}")
        
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        try:
            tmp.write(pdf_bytes)
            tmp.flush()
            tmp.close()
            return send_file(tmp.name, as_attachment=True, download_name=filename, mimetype="application/pdf")
        finally:
            try:
                os.unlink(tmp.name)
            except Exception:
                pass
    except Exception as exc:
        logger.error(f"Resume generation error: {str(exc)}")
        return _error(str(exc), 500)


@app.errorhandler(404)
def not_found(error):
    logger.warning(f"404 Error: {request.path}")
    return jsonify({"error": "Endpoint not found"}), 404


@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 Error: {str(error)}")
    return jsonify({"error": "Internal server error"}), 500


if __name__ == "__main__":
    debug_mode = os.environ.get('FLASK_ENV') == 'development'
    app.run(debug=debug_mode, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
