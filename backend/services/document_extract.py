"""Extract plain text from resume PDF/DOCX files."""

import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        raise RuntimeError("PyPDF2 not installed. Run: python -m pip install PyPDF2")

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read PDF: {exc}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("python-docx not installed. Run: python -m pip install python-docx")

    try:
        bio = io.BytesIO(file_bytes)
        doc = Document(bio)
        parts: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        # also try to extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX: {exc}")